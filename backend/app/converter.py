import os
import fitz
import json
import shutil
import logging
import subprocess
import platform
import pytesseract
from PIL import Image
from redis import Redis
from docx import Document
from pdf2docx import Converter
from reportlab.pdfgen import canvas
from app.aws_utils import upload_to_s3
from docx2pdf import convert as docx2pdf_convert

logger = logging.getLogger(__name__)
redis_conn = Redis()

# --- UTILITY HELPERS ---
def ensure_output_folder(path):
    os.makedirs(os.path.dirname(path), exist_ok=True)

# --- CONVERTERS ---
def convert_txt_to_pdf(input_path, output_path):
    ensure_output_folder(output_path)
    c = canvas.Canvas(output_path)
    with open(input_path, "r", encoding="utf-8") as f:
        text = f.read()
    for i, line in enumerate(text.splitlines()[:60]):
        c.drawString(50, 800 - (15 * i), line)
    c.save()
    return output_path


def convert_docx_to_pdf_cross_platform(input_path, output_path):
    system = platform.system().lower()

    if system == "windows":
        temp_path = output_path.replace(".pdf", ".docx")
        shutil.copyfile(input_path, temp_path)
        docx2pdf_convert(temp_path, output_path)
        os.remove(temp_path)
        return output_path

    else:
        try:
            subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.path.dirname(output_path),
                input_path
            ], check=True)

            converted = os.path.join(
                os.path.dirname(output_path),
                os.path.splitext(os.path.basename(input_path))[0] + ".pdf"
            )

            if not os.path.exists(converted):
                raise Exception("LibreOffice conversion failed: output file not found.")

            os.rename(converted, output_path)
            return output_path

        except Exception as e:
            raise Exception(f"LibreOffice conversion failed: {e}")


def convert_doc_to_image(input_path, output_path):
    ensure_output_folder(output_path)

    # Extract text from the DOCX file
    doc = Document(input_path)
    text = "\n".join([p.text for p in doc.paragraphs])

    # Create a temporary PDF file from the text
    temp_pdf = output_path.rsplit(".", 1)[0] + "_temp.pdf"
    c = canvas.Canvas(temp_pdf)
    for i, line in enumerate(text.splitlines()[:60]):
        c.drawString(50, 800 - (15 * i), line)
    c.save()

    # Convert the PDF to an image
    pdf_doc = fitz.open(temp_pdf)
    pix = pdf_doc[0].get_pixmap(dpi=300)
    pdf_doc.close()
    os.remove(temp_pdf)

    # Determine output image format
    output_ext = os.path.splitext(output_path)[1].lower().strip(".")
    if output_ext not in ["jpg", "jpeg", "png"]:
        raise ValueError(f"Unsupported image format: .{output_ext}")

    # Convert to PIL Image and save with correct format
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    image_format = "JPEG" if output_ext in ["jpg", "jpeg"] else output_ext.upper()
    img.save(output_path, format=image_format)

    return output_path


def convert_pdf_to_image(input_path, output_path):
    ensure_output_folder(output_path)
    pdf_doc = fitz.open(input_path)
    pix = pdf_doc[0].get_pixmap()
    pix.save(output_path)
    pdf_doc.close()
    return output_path


def convert_pdf_to_docx(input_path, output_path):
    logger.info(f"Converting PDF to DOCX: {input_path} → {output_path}")
    ensure_output_folder(output_path)

    try:
        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()

        if not os.path.exists(output_path) or os.path.getsize(output_path) < 1024:
            raise Exception("pdf2docx output too small or empty.")
        return output_path

    except Exception as e:
        logger.warning(f"pdf2docx failed: {e}. Trying OCR fallback...")

        pdf = fitz.open(input_path)
        full_text = ""
        for page in pdf:
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            full_text += pytesseract.image_to_string(img) + "\n"

        doc = Document()
        doc.add_paragraph(full_text.strip())
        doc.save(output_path)

        if not os.path.exists(output_path) or os.path.getsize(output_path) < 1024:
            raise Exception("OCR fallback also failed: empty output.")

        return output_path


def convert_image(input_path, output_path, output_format):
    ensure_output_folder(output_path)
    image = Image.open(input_path)
    if output_format == "pdf":
        image.convert("RGB").save(output_path)
    else:
        image.save(output_path)
    return output_path


def fallback_copy(input_path, output_path):
    ensure_output_folder(output_path)
    shutil.copyfile(input_path, output_path)
    return output_path


# --- MAIN DISPATCH FUNCTION ---
def convert_file(input_path, output_format, task_id):
    file_name = os.path.basename(input_path)
    input_ext = os.path.splitext(file_name)[1].lower()
    output_path = f"converted/{task_id}.{output_format}"

    try:
        logger.info(f"[{task_id}] Starting conversion: {file_name} to {output_format}")

        # DISPATCH ROUTES
        if input_ext in [".txt", ".rtf"] and output_format == "pdf":
            output_path = convert_txt_to_pdf(input_path, output_path)

        elif input_ext in [".docx", ".doc"]:
            if output_format == "pdf":
                output_path = convert_docx_to_pdf_cross_platform(input_path, output_path)
            elif output_format in ["png", "jpg", "jpeg"]:
                output_path = convert_doc_to_image(input_path, output_path)
            else:
                output_path = fallback_copy(input_path, output_path)

        elif input_ext == ".pdf":
            if output_format == "docx":
                output_path = convert_pdf_to_docx(input_path, output_path)
            elif output_format in ["jpg", "jpeg", "png"]:
                output_path = convert_pdf_to_image(input_path, output_path)
            else:
                output_path = fallback_copy(input_path, output_path)

        elif input_ext in [".jpg", ".jpeg", ".png"]:
            output_path = convert_image(input_path, output_path, output_format)

        else:
            output_path = fallback_copy(input_path, output_path)

        # Validate output file
        if not os.path.exists(output_path):
            raise Exception(f"Conversion failed: {output_path} was not created.")

        # Upload to S3
        s3_key = f"{task_id}.{output_format}"
        if not upload_to_s3(output_path, s3_key):
            raise Exception("Upload to S3 failed")

        # Redis update
        task_data = {
            "status": "completed",
            "output_format": output_format,
            "s3_key": s3_key
        }
        redis_conn.set(task_id, json.dumps(task_data))
        logger.info(f"[{task_id}] Conversion completed and uploaded.")

    except Exception as e:
        logger.error(f"[{task_id}] Conversion failed: {e}")
        task_data = {
            "status": "error",
            "error": str(e),
            "output_format": output_format
        }
        redis_conn.set(task_id, json.dumps(task_data))

    finally:
        try:
            if os.path.exists(input_path):
                os.remove(input_path)
                logger.info(f"[{task_id}] Cleaned up input file.")
            if os.path.exists(output_path):
                os.remove(output_path)
                logger.info(f"[{task_id}] Cleaned up output file.")
        except Exception as cleanup_error:
            logger.error(f"[{task_id}] Cleanup error: {cleanup_error}")

